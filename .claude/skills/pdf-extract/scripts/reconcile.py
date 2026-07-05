#!/usr/bin/env python3
"""Deterministic QA helpers for PDF -> Markdown extraction (pdf-extract skill).

A born-digital PDF's embedded text layer is the character-level ground truth:
no OCR/VLM output may disagree with it. `check` verifies that every text-layer
line appears, character-exact (after normalization), somewhere in a converted
output (.md or .docx), and reports each miss with its closest fuzzy match so an
agent can distinguish a misread (high ratio, a few chars differ) from dropped
content (low ratio). `probe` classifies pages digital vs raster (raster pages
have no oracle, so QA falls back to visual sampling). `pdfcheck` hygiene-checks
a searchable PDF's invisible text layer (extractability, NUL/control chars,
Simplified leaks). `render` exports a page PNG for visual QA.

Normalization on BOTH sides: NFKC (fullwidth/halfwidth), all whitespace removed
(CJK lines re-wrap freely in markdown), markdown syntax stripped, and optional
OpenCC s2tw when --locale zh-TW (the service converts its output, so the layer
must be compared in the same script).

Needs: pymupdf. Optional: opencc. Run with the repo's worker/.venv/bin/python
(has both) or any python with pymupdf installed.
"""

from __future__ import annotations

import argparse
import difflib
import json
import re
import unicodedata
from collections import Counter
from collections.abc import Callable
from pathlib import Path

import pymupdf as fitz

_S2TW: Callable[[str], str] | None
try:
    from opencc import OpenCC

    _S2TW = OpenCC("s2tw").convert
except Exception:  # optional dep; zh-TW comparison degrades gracefully
    _S2TW = None

_WS = re.compile(r"\s+")
_MD_IMG = re.compile(r"!\[[^\]]*\]\([^)]*\)")
_MD_LINK = re.compile(r"\[([^\]]*)\]\([^)]*\)")
_HTML_TAG = re.compile(r"</?[a-zA-Z][^>]*>")
_TABLE_SEP = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$", re.MULTILINE)
# Markdown / decoration chars removed from BOTH sides so markup never breaks a match.
# Bullet glyphs too: the VL renders •/◼/⚫/◦/▪ as ■, ☑, ·, or 。 unpredictably.
_MARKS = str.maketrans("", "", "|#*`>_\\~•·◦▪‣◼⚫●■☑✓☐○◆◇")
# TOC dot-leaders: runs of 2+ leader chars collapse to nothing on both sides (the layer
# and the VL rarely agree on how many dots a leader has).
_LEADER = re.compile(r"[.·・…]{2,}")
# Standalone page numbers, arabic or roman ("- 3 -", "– iv –").
_PAGE_NO = re.compile(r"^[\s\-—–~〜.·]*(\d{1,4}|[ivxlcdm]{1,7})[\s\-—–~〜.·]*$", re.IGNORECASE)


def norm(s: str, locale: str | None) -> str:
    s = unicodedata.normalize("NFKC", s)
    if locale == "zh-TW" and _S2TW:
        s = _S2TW(s)
    s = _LEADER.sub("", s)
    return _WS.sub("", s).translate(_MARKS)


def _docx_text(path: Path) -> str:
    """All document text from a .docx: body paragraphs + every table cell."""
    import docx  # python-docx (in the worker venv)

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def out_corpus(path: Path, locale: str | None) -> str:
    """Normalized text corpus of a converted output (.md or .docx)."""
    if path.suffix.lower() == ".docx":
        return norm(_docx_text(path), locale)
    text = path.read_text(encoding="utf-8")
    text = _MD_IMG.sub("", text)
    text = _MD_LINK.sub(r"\1", text)
    text = _TABLE_SEP.sub("", text)
    text = _HTML_TAG.sub("", text)
    return norm(text, locale)


def fuzzy(needle: str, hay: str) -> tuple[str, float]:
    """Closest window of `hay` to `needle` (locates a misread's counterpart)."""
    sm = difflib.SequenceMatcher(None, needle, hay, autojunk=False)
    m = sm.find_longest_match(0, len(needle), 0, len(hay))
    lo = max(0, m.b - m.a)
    win = hay[lo : min(len(hay), lo + len(needle) + 4)]
    return win, round(difflib.SequenceMatcher(None, needle, win).ratio(), 3)


def cmd_probe(args: argparse.Namespace) -> None:
    doc = fitz.open(args.pdf)
    pages = []
    for i in range(doc.page_count):
        page = doc.load_page(i)
        chars = len(_WS.sub("", page.get_text()))
        pages.append(
            {
                "page": i + 1,
                "text_chars": chars,
                "images": len(page.get_images(full=True)),
                "kind": "digital" if chars >= args.digital_chars else "raster",
            }
        )
    dig = sum(1 for p in pages if p["kind"] == "digital")
    cls = "digital" if dig == len(pages) else ("raster" if dig == 0 else "mixed")
    print(
        json.dumps(
            {
                "pdf": args.pdf,
                "pages": len(pages),
                "digital_pages": dig,
                "raster_pages": len(pages) - dig,
                "classification": cls,
                "detail": pages,
            },
            ensure_ascii=False,
        )
    )


def cmd_check(args: argparse.Namespace) -> None:
    locale = args.locale
    doc = fitz.open(args.pdf)
    hay = out_corpus(Path(args.out), locale)
    if locale == "zh-TW" and _S2TW is None:
        print("WARNING: opencc not installed; zh-TW comparison without s2tw may false-flag", flush=True)

    # A line repeating verbatim on >=3 pages is a header/footer/watermark or a repeated
    # table header (which the service's table merging dedupes on purpose), never body
    # prose — so its absence from the markdown is reported separately, not as a miss.
    page_lines: list[list[str]] = []
    freq: Counter[str] = Counter()
    for i in range(doc.page_count):
        lines = [ln.strip() for ln in doc.load_page(i).get_text().splitlines() if ln.strip()]
        page_lines.append(lines)
        freq.update({norm(ln, locale) for ln in lines if norm(ln, locale)})
    repeated = {k for k, v in freq.items() if v >= 3}

    pages_out: list[dict[str, object]] = []
    total_chars = covered_chars = 0
    n_missing = pages_with_misses = 0
    for pno, lines in enumerate(page_lines, 1):
        missing: list[dict[str, object]] = []
        ignorable_missing: list[str] = []
        for raw in lines:
            key = norm(raw, locale)
            if len(key) < args.min_len:
                continue
            ignorable = key in repeated or bool(_PAGE_NO.match(raw))
            if key in hay:
                if not ignorable:
                    total_chars += len(key)
                    covered_chars += len(key)
                continue
            if ignorable:
                ignorable_missing.append(raw)
                continue
            total_chars += len(key)
            entry: dict[str, object] = {"text": raw}
            if len(missing) < args.max_fuzzy:  # fuzzy is the costly part; cap per page
                win, ratio = fuzzy(key, hay)
                entry["closest_in_md"] = win[:160]
                entry["ratio"] = ratio
            missing.append(entry)
        if missing or ignorable_missing:
            pages_out.append(
                {"page": pno, "missing": missing, "ignorable_missing": ignorable_missing}
            )
            n_missing += len(missing)
            pages_with_misses += 1 if missing else 0

    print(
        json.dumps(
            {
                "pdf": args.pdf,
                "out": args.out,
                "locale": locale,
                "summary": {
                    "pages": len(page_lines),
                    "pages_with_misses": pages_with_misses,
                    "missing_lines": n_missing,
                    "char_coverage": round(covered_chars / total_chars, 4) if total_chars else None,
                    "clean": n_missing == 0,
                },
                "pages": pages_out,
            },
            ensure_ascii=False,
            indent=1,
        )
    )


def cmd_pdfcheck(args: argparse.Namespace) -> None:
    """Hygiene check on a searchable PDF's (invisible) text layer."""
    doc = fitz.open(args.pdf)
    pages = []
    nuls = ctrl = 0
    simplified: list[dict[str, object]] = []
    for i in range(doc.page_count):
        text = doc.load_page(i).get_text()
        nuls += text.count("\x00")
        ctrl += sum(1 for ch in text if ord(ch) < 32 and ch not in "\n\t\r")
        if args.locale == "zh-TW" and _S2TW:
            for ln in text.splitlines():
                ln = ln.strip()
                if ln and _S2TW(ln) != ln and len(simplified) < 20:
                    simplified.append({"page": i + 1, "text": ln, "s2tw": _S2TW(ln)})
        pages.append({"page": i + 1, "text_chars": len(_WS.sub("", text))})
    empty = [p["page"] for p in pages if int(str(p["text_chars"])) == 0]
    print(
        json.dumps(
            {
                "pdf": args.pdf,
                "summary": {
                    "pages": len(pages),
                    "empty_text_pages": empty,
                    "nul_chars": nuls,
                    "control_chars": ctrl,
                    "simplified_leaks": len(simplified),
                },
                "simplified_samples": simplified,
                "pages": pages,
            },
            ensure_ascii=False,
            indent=1,
        )
    )


def cmd_render(args: argparse.Namespace) -> None:
    doc = fitz.open(args.pdf)
    page = doc[args.page - 1]
    pix = page.get_pixmap(matrix=fitz.Matrix(args.zoom, args.zoom))
    pix.save(args.out)
    print(json.dumps({"out": args.out, "page": args.page, "size": f"{pix.width}x{pix.height}"}))


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    sub = ap.add_subparsers(dest="cmd", required=True)

    p = sub.add_parser("probe", help="classify pages digital vs raster")
    p.add_argument("pdf")
    p.add_argument("--digital-chars", type=int, default=50, help="min text chars for a digital page")
    p.set_defaults(fn=cmd_probe)

    c = sub.add_parser("check", help="reconcile a converted output (.md or .docx) against the PDF text layer")
    c.add_argument("pdf")
    c.add_argument("out")
    c.add_argument("--locale", default=None, help="zh-TW applies s2tw to both sides")
    c.add_argument("--min-len", type=int, default=2, help="ignore normalized lines shorter than this")
    c.add_argument("--max-fuzzy", type=int, default=30, help="fuzzy-match at most this many misses per page")
    c.set_defaults(fn=cmd_check)

    k = sub.add_parser("pdfcheck", help="hygiene check on a searchable PDF's text layer")
    k.add_argument("pdf")
    k.add_argument("--locale", default=None, help="zh-TW also detects Simplified leaks")
    k.set_defaults(fn=cmd_pdfcheck)

    r = sub.add_parser("render", help="render one page to PNG for visual QA")
    r.add_argument("pdf")
    r.add_argument("page", type=int)
    r.add_argument("out")
    r.add_argument("--zoom", type=float, default=2.0)
    r.set_defaults(fn=cmd_render)

    args = ap.parse_args()
    args.fn(args)


if __name__ == "__main__":
    main()
