"""Markdown/HTML/DOCX post-processing shared by the doc-parse pipelines.

Ports the Flask app's HTML-table/image -> Markdown conversion and Word style
fix-ups, plus multi-page DOCX merging via docxcompose (for native save_to_word).
"""

from __future__ import annotations

import base64
import contextlib
import html
import logging
import re
import uuid
from collections.abc import Callable
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup, Tag

from .text_utils import fix_ocr_text, latex_to_unicode, strip_cjk_spaces, strip_cjk_spaces_across

log = logging.getLogger("worker.postprocess")

_TABLE_PATTERNS = [
    re.compile(
        r"<div[^>]*>\s*<html>\s*<body>\s*<table[^>]*>.*?</table>\s*</body>\s*</html>\s*</div>",
        re.DOTALL,
    ),
    re.compile(r'<div[^>]*class="[^"]*table[^"]*"[^>]*>.*?</div>', re.DOTALL),
    re.compile(r"<table[^>]*>.*?</table>", re.DOTALL),
]
_IMG_PATTERNS = [
    re.compile(r"<div[^>]*>\s*<img[^>]+/?>\s*</div>", re.DOTALL),
    re.compile(r"<img[^>]+/?>", re.DOTALL),
]
_EXT_BY_MEDIA = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/gif": ".gif",
    "image/webp": ".webp",
    "image/svg+xml": ".svg",
}
# PaddleOCR-VL wraps figure/table/chart captions in a centered div; markdown has no div
# semantics, so unwrap it (keep inner content). Applied AFTER image extraction.
_CENTER_DIV = re.compile(r'<div style="text-align: center;">(.*?)</div>', re.DOTALL)

# VL's HTML-table OCR text can carry a literal two-char backslash-n ("\n") inside a cell (an
# enumerated list within one <td>); html_table_to_markdown's .split() collapses whitespace but
# NOT that literal sequence, so it survives into the GFM cell. In a native GFM pipe row a raw
# "\n" renders verbatim (GitHub/VS Code/most renderers) — the in-cell line break must be "<br>".
# A pipe-table BODY row is a line that both starts and ends with "|" (outside code fences); a
# separator row is all dashes/colons/spaces/pipes. We only rewrite "\n" on such rows, so a
# literal backslash-n in prose (a Windows path "C:\name", a code sample) is never touched.
_LITERAL_NL = re.compile(r"\\n")
_TABLE_SEP_ROW = re.compile(r"^\|[\s:|-]+\|$")


# Upper bound on a single cell's colspan/rowspan. A real document table never spans anywhere
# near this many rows/cols, but VL can emit a garbage span (e.g. "9999"); expanding it verbatim
# balloons every row's width (blank cells in Markdown, empty grid cells in DOCX) and can blow up
# memory. Clamp to keep normal tables untouched while bounding the pathological case.
_MAX_SPAN = 64


def _cell_span(cell: Any, attr: str) -> int:
    # VL can emit noisy/annotated spans (e.g. "2位", "2.0"); a bare int() would raise and
    # fail the whole export. Take the leading integer, clamp to [1, _MAX_SPAN], default 1.
    m = re.match(r"\d+", str(cell.get(attr) or "").strip())
    return min(_MAX_SPAN, max(1, int(m.group()))) if m else 1


def _carry_walk_rows[CellT](
    trs: list[Any],
    extract: Callable[[Tag], CellT],
    blank: Callable[[], CellT],
) -> list[list[CellT]]:
    """Shared rowspan/colspan carry-walk for both the Markdown and DOCX table paths.

    The cell PAYLOAD type (`CellT`) is a Markdown text string, or a list of DOCX (kind, value)
    segments; it is invariant across the walk and supplied by each caller.

    HTML tables allow row/col spans; both consumers need a rectangular grid instead. This walks
    each <tr>'s cells in order, expanding a colspan into (content, blank, blank...) and carrying a
    rowspan downward as a blank in the held column so continuation rows stay column-aligned. Only
    the cell PAYLOAD differs between callers (a Markdown text string vs. a list of DOCX segments),
    so it is supplied via `extract`; `blank` is a FACTORY (not a shared value) so each empty cell
    gets its own object — the DOCX blank is a mutable list the downstream writer may touch. Both
    callers proved byte-identical to their former inline copies of this walk (test_postprocess.py)."""
    out_rows: list[list[CellT]] = []
    carry: dict[int, int] = {}  # column -> remaining rowspan-continuation rows to leave blank
    for tr in trs:
        if not isinstance(tr, Tag):
            continue
        row: list[CellT] = []
        col = 0
        for cell in tr.find_all(["td", "th"]):
            if not isinstance(cell, Tag):
                continue
            while col in carry:  # columns held by an earlier row's rowspan -> blank
                row.append(blank())
                carry[col] -= 1
                if carry[col] <= 0:
                    del carry[col]
                col += 1
            payload = extract(cell)
            colspan = _cell_span(cell, "colspan")
            rowspan = _cell_span(cell, "rowspan")
            for k in range(colspan):
                row.append(payload if k == 0 else blank())  # content in the first cell, blanks after
                if rowspan > 1:
                    carry[col] = rowspan - 1
                col += 1
        while col in carry:  # rowspan columns trailing past this row's own cells
            row.append(blank())
            carry[col] -= 1
            if carry[col] <= 0:
                del carry[col]
            col += 1
        out_rows.append(row)
    return out_rows


def html_table_to_markdown(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    table = soup.find("table")
    if not table or not isinstance(table, Tag):
        return html
    rows = table.find_all("tr")
    if not rows:
        return html

    def _cell_text(cell: Tag) -> str:
        for br in cell.find_all("br"):
            br.replace_with(" ")
        text = " ".join(cell.get_text(separator=" ", strip=True).replace("|", "\\|").split())
        # Empty text collapses to the blank sentinel so an empty leading colspan cell matches the
        # trailing blanks (the former inline code's `text if (i == 0 and text) else " "`).
        return text if text else " "

    # Drop rows the walk produced empty (former code skipped `if row_cells`), then pad to width.
    md_rows = [r for r in _carry_walk_rows(rows, _cell_text, lambda: " ") if r]
    if not md_rows:
        return html
    max_cols = max(len(r) for r in md_rows)
    for r in md_rows:
        r.extend([" "] * (max_cols - len(r)))
    out = ["| " + " | ".join(md_rows[0]) + " |", "| " + " | ".join(["---"] * max_cols) + " |"]
    out.extend("| " + " | ".join(r) + " |" for r in md_rows[1:])
    return "\n\n" + "\n".join(out) + "\n\n"


# --- DOCX merged-cell (rowspan/colspan) support ------------------------------------------------
# PaddleX's writer consumes only the rectangular grid word_table_grid returns; it never emits a
# real w:vMerge/w:gridSpan, so a merged region ships as separate blank bordered boxes with divider
# lines drawn through it. We CANNOT touch that writer, so word_table_grid records, per table, the
# spans it flattened; save_word drains this collector and post-processes the final .docx, walking
# <w:tbl> IN THE SAME ORDER the grids were produced and applying the OOXML merge markup. The
# collector is module-level (the monkeypatched parser has no other channel back to save_word) and
# is reset at the start of each save_word run, so a stale run's specs never leak into the next.
#
# A MergeSpec is one per table: {(row, col): (rowspan, colspan)} keyed by each merge ORIGIN's grid
# coordinate (only origins with a span > 1 are recorded; unmerged cells are absent). row/col are
# 0-based indices into the rectangular grid word_table_grid returned. This is enough to reconstruct
# both directions: gridSpan on the origin (+ delete the colspan-1 blank cells in every row it
# covers) and vMerge restart on the origin with continue on the rows below.
#
# Each collector entry carries the spec PLUS a structural fingerprint (row count, first-row cell
# count, first-row text): word_table_grid only sees HTML tables, but PaddleX's writer can emit
# additional <w:tbl> for non-HTML content, so specs can NOT be zipped positionally against the
# body's tables — each spec must find ITS table by fingerprint (in order), and anything that
# doesn't match is left untouched.
MergeSpec = dict[tuple[int, int], tuple[int, int]]
TableSpecEntry = tuple[MergeSpec, int, int, str]  # (spec, n_rows, n_first_row_cells, first_row_fp)
_WORD_TABLE_MERGE_SPECS: list[TableSpecEntry] = []


def reset_word_table_merge_specs() -> None:
    """Clear the per-table span collector before a save_word run (guards against stale leakage)."""
    _WORD_TABLE_MERGE_SPECS.clear()


def drain_word_table_merge_specs() -> list[TableSpecEntry]:
    """Return and clear the span specs word_table_grid recorded (one per table, in call order)."""
    specs = list(_WORD_TABLE_MERGE_SPECS)
    _WORD_TABLE_MERGE_SPECS.clear()
    return specs


def word_table_grid(html: str) -> list[list[list[tuple[str, str]]]]:
    """Rectangular cell grid for PaddleX's DOCX table writer, honoring rowspan/colspan.

    PaddleX's own `_parse_html_table` (word_converter.py — unfixed in 3.7.2 and on upstream
    main) reads each <tr>'s <td>/<th> in order and IGNORES rowspan/colspan, then pads short rows
    at the END. So a vertically-merged cell's continuation row shifts one column left with a
    blank TRAILING cell (its Markdown export is correct; only the DOCX is wrong). We monkeypatch
    this in as the parser (docparse._patch_paddlex_word_table) so content lands in the right
    column: a rowspan carries a blank cell downward, a colspan expands into blanks. Returns the
    exact shape the writer consumes — rows of cells, each cell a list of (kind, value) segments
    with kind in {"text", "img"}; a spanned/blank cell is an empty list.

    ALSO records this table's span origins on _WORD_TABLE_MERGE_SPECS (drained by save_word) so the
    final-docx pass can replace those flattened blank cells with real w:vMerge/w:gridSpan markup.
    """
    from bs4.element import NavigableString

    def _cell_segments(cell: Tag) -> list[tuple[str, str]]:
        # Segment extraction ported verbatim from PaddleX's _parse_html_table.
        segments: list[tuple[str, str]] = []
        for child in cell.children:
            if isinstance(child, NavigableString):
                text = child.strip()
                if text:
                    segments.append(("text", text))
            elif isinstance(child, Tag) and child.name == "img":
                src = str(child.get("src", "") or "")
                if src:
                    segments.append(("img", src))
            elif isinstance(child, Tag):
                text = child.get_text(strip=True)
                if text:
                    segments.append(("text", text))
        return segments

    soup = BeautifulSoup(html, "html.parser")
    trs = soup.find_all("tr")
    # A spanned/blank cell is a fresh empty segment list; every row is kept (the writer wants the
    # full grid, unlike the Markdown path which drops empties).
    grid = _carry_walk_rows(trs, _cell_segments, lambda: [])
    # Fingerprint from the grid the writer will actually emit: row count, first-row cell count,
    # and the first row's text (normalized) — used to locate THIS table among the final body's
    # <w:tbl> elements, which can include non-HTML tables this collector never saw.
    fp = _row_fingerprint(
        "".join(v for cell in grid[0] for k, v in cell if k == "text") if grid else ""
    )
    _WORD_TABLE_MERGE_SPECS.append(
        (_word_table_spans(trs), len(grid), len(grid[0]) if grid else 0, fp)
    )
    return grid


def _row_fingerprint(text: str) -> str:
    """Whitespace-normalized row text for spec<->table matching."""
    return "".join(text.split())


def _word_table_spans(trs: list[Any]) -> MergeSpec:
    """Compute {(row, col): (rowspan, colspan)} for span ORIGINS, mirroring _carry_walk_rows.

    Runs the identical carry-walk so the (row, col) it emits index EXACTLY the same rectangular
    grid word_table_grid returned — only cells whose rowspan or colspan exceeds 1 are recorded.
    Any span is clamped to [1, _MAX_SPAN] like the grid, and bounded to the rows/cols actually
    present so a garbage rowspan can't point past the table.
    """
    spans: MergeSpec = {}
    carry: dict[int, int] = {}
    row_idx = 0
    for tr in trs:
        if not isinstance(tr, Tag):
            continue
        col = 0
        for cell in tr.find_all(["td", "th"]):
            if not isinstance(cell, Tag):
                continue
            while col in carry:  # skip columns held by an earlier row's rowspan
                carry[col] -= 1
                if carry[col] <= 0:
                    del carry[col]
                col += 1
            colspan = _cell_span(cell, "colspan")
            rowspan = _cell_span(cell, "rowspan")
            if rowspan > 1 or colspan > 1:
                spans[(row_idx, col)] = (rowspan, colspan)
            if rowspan > 1:
                for k in range(colspan):
                    carry[col + k] = rowspan - 1
            col += colspan
        while col in carry:  # drain trailing rowspan-held columns
            carry[col] -= 1
            if carry[col] <= 0:
                del carry[col]
            col += 1
        row_idx += 1
    return spans


def _apply_table_merges(tbl: Any, spec: MergeSpec) -> None:
    """Apply w:gridSpan / w:vMerge to one <w:tbl> from its recorded span origins (in place).

    OOXML rules: a horizontal merge sets w:gridSpan=N on the origin cell and REMOVES the N-1
    cells to its right in that row; a vertical merge sets w:vMerge=restart on the origin and
    w:vMerge (continue) on the cell directly below in each spanned row (that cell is KEPT). A
    combined span does both: every row the merge covers first collapses its colspan (gridSpan +
    delete), then the continuation rows' surviving merged cell gets a vMerge continue.

    Robustness: if the table's realized shape doesn't match the spec (merge_docx reordering, a
    writer that dropped/added rows), any out-of-range coordinate is skipped, so a mismatch degrades
    to "no merge applied" rather than corrupting the table. Tables with an empty spec are untouched
    (byte-identical), which every no-span table proves via the existing writer tests.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not spec:
        return
    # Snapshot each row's cells as a list up front; we mutate <w:tr> children as we go and must not
    # let python-docx re-scan a half-edited tree. Index by the ORIGINAL grid column (pre-deletion).
    rows = tbl.findall(qn("w:tr"))
    grid_rows = [r.findall(qn("w:tc")) for r in rows]

    def _tc_pr(tc: Any) -> Any:
        pr = tc.find(qn("w:tcPr"))
        if pr is None:
            pr = OxmlElement("w:tcPr")
            tc.insert(0, pr)  # tcPr must be the first child of tc
        return pr

    def _set_gridspan(tc: Any, n: int) -> None:
        pr = _tc_pr(tc)
        for old in pr.findall(qn("w:gridSpan")):
            pr.remove(old)
        gs = OxmlElement("w:gridSpan")
        gs.set(qn("w:val"), str(n))
        pr.insert(0, gs)

    def _set_vmerge(tc: Any, restart: bool) -> None:
        pr = _tc_pr(tc)
        for old in pr.findall(qn("w:vMerge")):
            pr.remove(old)
        vm = OxmlElement("w:vMerge")
        if restart:
            vm.set(qn("w:val"), "restart")
        pr.append(vm)

    for (r0, c0), (rowspan, colspan) in spec.items():
        if r0 >= len(grid_rows) or c0 >= len(grid_rows[r0]):
            continue  # shape mismatch -> skip this merge, never corrupt
        r1 = min(r0 + rowspan, len(grid_rows))
        # Horizontal: collapse the colspan in EVERY row the (possibly vertical) merge covers, so
        # each row keeps one merged cell in that column. Delete right-to-left to preserve indices.
        if colspan > 1:
            for r in range(r0, r1):
                cells = grid_rows[r]
                if c0 >= len(cells):
                    continue
                span = min(colspan, len(cells) - c0)
                _set_gridspan(cells[c0], span)
                for c in range(c0 + span - 1, c0, -1):
                    rows[r].remove(cells[c])
        # Vertical: origin restart, each row below continues on the (now merged) column-c0 cell.
        if rowspan > 1:
            _set_vmerge(grid_rows[r0][c0], restart=True)
            for r in range(r0 + 1, r1):
                if c0 < len(grid_rows[r]):
                    _set_vmerge(grid_rows[r][c0], restart=False)


def html_img_to_markdown(html: str, output_dir: Path | None) -> str:
    soup = BeautifulSoup(html, "html.parser")
    img = soup.find("img")
    if not img:
        return html
    src = str(img.get("src") or "")
    alt = str(img.get("alt") or "")
    if not src:
        return html
    if alt.strip().lower() == "image":
        alt = ""
    alt = alt.replace("\\", "\\\\").replace("[", "\\[").replace("]", "\\]")
    if src.startswith("data:"):
        if not (output_dir and output_dir.is_dir()):
            return f"\n\n<!-- Image: {alt or 'embedded'} (data URI) -->\n\n"
        m = re.match(r"data:([^;,]+)?(?:;base64)?,(.+)", src, re.DOTALL)
        if not m:
            return f"\n\n<!-- Image: {alt or 'embedded'} (invalid data URI) -->\n\n"
        ext = _EXT_BY_MEDIA.get(m.group(1) or "", ".bin")
        name = f"extracted_{uuid.uuid4().hex[:8]}{ext}"
        try:
            (output_dir / name).write_bytes(base64.b64decode(m.group(2)))
            src = name
        except Exception:
            return f"\n\n<!-- Image: {alt or 'embedded'} (extract failed) -->\n\n"
    src = src.replace(" ", "%20").replace("(", "%28").replace(")", "%29")
    return f"\n\n![{alt}]({src})\n\n"


def _replace_all(text: str, patterns: list[re.Pattern[str]], fn: Callable[[str], str]) -> str:
    # Convert every match in one pass per pattern (unbounded; earlier patterns are more
    # specific, so they run first). fn receives the full matched chunk (m.group(0)).
    for pat in patterns:
        text = pat.sub(lambda m: fn(m.group(0)), text)
    return text


def _fix_pipe_table_cell_newlines(md: str) -> str:
    """Turn a literal two-char "\\n" into a GFM "<br>" inside native pipe-table body rows.

    VL's HTML-table cells sometimes hold an enumerated list joined by a literal backslash-n; after
    html_table_to_markdown that ships raw in the GFM cell (and renders verbatim). Only lines that
    are pipe-table body rows (start+end with "|", not a `| --- |` separator, not inside a ``` fence
    or an INDENTED code block — GFM treats >=4 leading spaces as code, and a table row may be
    indented at most 3) are rewritten, so a literal "\\n" in ordinary prose/paths stays untouched.
    Idempotent: once a row's "\\n" is "<br>", there is no literal backslash-n left to match.
    """
    if "\\n" not in md:  # cheap fast-path: no literal backslash-n anywhere
        return md
    in_fence = False
    out: list[str] = []
    for line in md.split("\n"):
        stripped = line.strip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            in_fence = not in_fence
            out.append(line)
            continue
        indented_code = line.startswith("    ") or line.startswith("\t")
        is_row = (
            not in_fence
            and not indented_code
            and stripped.startswith("|")
            and stripped.endswith("|")
            and not _TABLE_SEP_ROW.match(stripped)
        )
        out.append(_LITERAL_NL.sub("<br>", line) if is_row else line)
    return "\n".join(out)


def process_markdown(md: str, output_dir: Path | None, locale: str | None, images: bool) -> str:
    """Fix S->T, strip stray CJK spaces, render inline LaTeX, convert HTML tables to md."""
    md = fix_ocr_text(md, locale)
    md = strip_cjk_spaces(md)  # before table/latex passes; only touches Han<->Han gaps, not "| "
    md = latex_to_unicode(md)
    md = _replace_all(md, _TABLE_PATTERNS, html_table_to_markdown)
    md = _fix_pipe_table_cell_newlines(md)  # literal "\n" in a GFM cell -> "<br>" (after HTML->GFM)
    if images:
        md = _replace_all(md, _IMG_PATTERNS, lambda chunk: html_img_to_markdown(chunk, output_dir))
    # Unwrap VL's presentational center-div (figure/table/chart caption blocks) left after
    # table + image conversion; loop because captions can nest. Runs last so any data-URI
    # <img> inside such a div was already extracted to disk by the image pass above.
    prev = None
    while prev != md:
        prev = md
        md = _CENTER_DIV.sub(lambda m: m.group(1), md)
    # Decode HTML entities (`&amp;`, `&#8451;`->℃, `&deg;`->°) that VL leaves in prose outside
    # tables — BS4 already decoded them inside the tables it parsed above, so this runs last
    # and only touches what remains, and it leaves bare `<`/`&`/`%` alone.
    md = html.unescape(md)
    # Re-strip Han<->Han spaces: the passes above can EXPOSE ones the early strip couldn't see
    # (an entity decoding to a space, markup between the two Han chars) — real 06-doc output
    # shipped "繳交 維護人員名冊" past the early strip. Idempotent, so a second pass is free.
    return strip_cjk_spaces(md)


def merge_docx(paths: list[Path], out_path: Path) -> None:
    """Merge multiple .docx into one (native save_to_word is per-result)."""
    from docx import Document

    if len(paths) == 1:
        paths[0].replace(out_path)
        return
    from docxcompose.composer import Composer

    master = Document(str(paths[0]))
    composer = Composer(master)
    for p in paths[1:]:
        composer.append(Document(str(p)))
    composer.save(str(out_path))


def fix_word_styles(docx_path: Path) -> None:
    """Force black text + add table borders (normalizes generator quirks)."""
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import RGBColor

        doc = Document(str(docx_path))
        black = RGBColor(0, 0, 0)
        for para in doc.paragraphs:
            for run in para.runs:
                with contextlib.suppress(Exception):
                    run.font.color.rgb = black
        for table in doc.tables:
            try:
                tbl = table._tbl
                tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
                if tbl.tblPr is None:
                    tbl.insert(0, tblPr)
                borders = OxmlElement("w:tblBorders")
                for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    el = OxmlElement(f"w:{name}")
                    el.set(qn("w:val"), "single")
                    el.set(qn("w:sz"), "4")
                    el.set(qn("w:space"), "0")
                    el.set(qn("w:color"), "000000")
                    borders.append(el)
                existing = tblPr.find(qn("w:tblBorders"))
                if existing is not None:
                    tblPr.remove(existing)
                # OOXML CT_TblPr requires a fixed child order; insert tblBorders before its
                # legal successors (append would place it after tblLook, which strict parsers
                # reject). insert_element_before falls back to append when none are present.
                tblPr.insert_element_before(
                    borders,
                    "w:tblLayout",
                    "w:tblCellMar",
                    "w:tblLook",
                    "w:tblCaption",
                    "w:tblDescription",
                    "w:tblPrChange",
                )
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                with contextlib.suppress(Exception):
                                    run.font.color.rgb = black
            except Exception as e:
                log.warning("table style fix skipped: %s", e)
        tmp = docx_path.with_suffix(".docx.tmp")
        doc.save(str(tmp))
        tmp.replace(docx_path)
    except Exception as e:
        log.warning("word style fix skipped: %s", e)


def apply_word_table_merges(docx_path: Path, specs: list[TableSpecEntry]) -> None:
    """Rewrite the final DOCX's tables with real w:vMerge/w:gridSpan from the drained span specs.

    `specs` are ordered as word_table_grid recorded them, but the body's <w:tbl> sequence can
    contain ADDITIONAL tables the collector never saw (PaddleX's writer also emits <w:tbl> for
    non-HTML table content), so positional zipping mis-targets. Each spec instead finds ITS
    table by fingerprint — row count, first-row cell count, first-row text — scanning forward
    from the previous match (both sequences are in body order). A spec whose table can't be
    found is skipped; a table matching no spec is untouched — mismatches degrade to "no merge",
    never corruption. All-empty specs leave the DOCX byte-identical.
    """
    if not any(spec for spec, _r, _c, _f in specs):  # no table had a rowspan/colspan
        return
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(docx_path))
        tbls = doc.element.body.findall(qn("w:tbl"))

        def _tbl_matches(tbl: Any, n_rows: int, n_cells: int, fp: str) -> bool:
            trs = tbl.findall(qn("w:tr"))
            if len(trs) != n_rows or not trs:
                return False
            cells = trs[0].findall(qn("w:tc"))
            if len(cells) != n_cells:
                return False
            first_row_text = "".join(t.text or "" for t in trs[0].iter(qn("w:t")))
            return _row_fingerprint(first_row_text) == fp

        ti = 0
        for spec, n_rows, n_cells, fp in specs:
            while ti < len(tbls):
                tbl = tbls[ti]
                ti += 1
                if _tbl_matches(tbl, n_rows, n_cells, fp):
                    if spec:
                        try:
                            _apply_table_merges(tbl, spec)
                        except Exception as e:
                            log.warning("table merge apply skipped: %s", e)
                    break
        tmp = docx_path.with_suffix(".docx.tmp")
        doc.save(str(tmp))
        tmp.replace(docx_path)
    except Exception as e:
        log.warning("word table merge pass skipped: %s", e)


# Heading detection for apply_word_heading_styles. PaddleX's writer discards each block's semantic
# `level`, styling headings only via font size+bold+alignment (build_word_blocks BASE_STYLE_MAP):
# a heading block is bold, size >= _HEADING_MIN_PT, and NOT justified (body text is size-12
# JUSTIFY). Ranking the distinct heading sizes DESCENDING maps the largest -> Heading 1, next ->
# Heading 2, ... (capped at Heading _MAX_HEADING_LEVELS). This recovers a navigable outline from
# the ONLY structural signal the writer preserved, and is guarded so it never fires on a document
# whose bold runs aren't really headings (see _MAX_HEADING_FRACTION).
_HEADING_MIN_PT = 13.5  # size-14 heading vs size-12 body; a little slack for rounding
_MAX_HEADING_LEVELS = 4
_MAX_HEADING_FRACTION = 0.5  # if > half the paragraphs look like headings, the signal is noise -> skip


def _heading_signature(para: Any) -> tuple[float, bool] | None:
    """(font_size_pt, is_left_or_center) for a heading-looking paragraph, else None.

    A heading is: non-empty text, bold FIRST run, size >= _HEADING_MIN_PT, and not JUSTIFY-aligned
    (body prose is size-12 justified). Returns the size so the caller can rank distinct sizes into
    Heading levels; a non-heading paragraph returns None.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not para.runs or not para.text.strip():
        return None
    run = para.runs[0]
    size = run.font.size
    if size is None or not run.font.bold:
        return None
    pt = size.pt
    if pt < _HEADING_MIN_PT:
        return None
    if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return None
    return (pt, True)


def apply_word_heading_styles(docx_path: Path, levels: dict[str, int] | None = None) -> None:
    """Promote size/bold heading paragraphs to real Word 'Heading N' styles (navigation outline).

    PaddleX ships every paragraph as 'Normal' (headings differ only in size/bold), so Word's
    navigation pane, cross-references and TOC see a flat document. We reconstruct the outline from
    the writer's own heading styling: collect the distinct font sizes of heading-looking paragraphs
    (bold, big, not justified), rank them DESCENDING to Heading 1..N, and set each qualifying
    paragraph's style accordingly.

    `levels` recovers heading DEPTH the size rank cannot: PaddleX styles EVERY heading level at
    the same ~14pt bold, so size-ranking alone flattens the outline to Heading 1+2. The same VL
    pass's markdown carries the true levels; the caller passes {whitespace-normalized heading
    text: markdown level}, and a text-matched paragraph takes that level (size rank stays the
    fallback for unmatched headings).

    Guards make it conservative and safe:
      * do nothing if NO paragraph qualifies, or if MORE than _MAX_HEADING_FRACTION of paragraphs
        qualify (then 'bold + big' isn't really a heading signal — bail rather than mis-style body);
      * the built-in 'Heading N' styles always exist in python-docx's template, so the assignment
        can't raise for a missing style; a per-paragraph failure is suppressed.
    Idempotent: a paragraph already on 'Heading N' keeps its run size/bold, so a second pass
    resolves to the same level.
    """
    try:
        from docx import Document

        doc = Document(str(docx_path))
        paras = [p for p in doc.paragraphs if p.text.strip()]
        if not paras:
            return
        sigs = [(p, _heading_signature(p)) for p in paras]
        heads = [(p, s) for p, s in sigs if s is not None]
        if not heads or len(heads) > len(paras) * _MAX_HEADING_FRACTION:
            return  # no headings, or so many that the signal is unreliable -> leave as-is
        sizes = sorted({s[0] for _, s in heads}, reverse=True)  # largest size = shallowest heading
        level_of = {sz: min(i + 1, _MAX_HEADING_LEVELS) for i, sz in enumerate(sizes)}
        for para, sig in heads:
            level = (levels or {}).get("".join(para.text.split()))
            if level is None:
                level = level_of[sig[0]]
            with contextlib.suppress(Exception):
                para.style = doc.styles[f"Heading {min(level, _MAX_HEADING_LEVELS)}"]
        tmp = docx_path.with_suffix(".docx.tmp")
        doc.save(str(tmp))
        tmp.replace(docx_path)
    except Exception as e:
        log.warning("word heading style pass skipped: %s", e)


def finalize_docx_text(docx_path: Path, locale: str | None) -> None:
    """Normalize the text of a generated DOCX in place.

    The native `save_to_word` path (PaddleOCR-VL) never routes through process_markdown, so
    without this its output ships as-is: (1) HTML entities (`&amp;`, `&#8451;`) as literal
    source — decoded; (2) inline LaTeX math/units (`$ ^{\\circ} $C`) as literal source —
    rendered to Unicode for every locale; (3) Simplified characters on a zh-TW job — converted
    to Traditional. (1)/(2) are local to a single run and run per-<w:t>. (3) is NOT: OpenCC s2tw
    is phrase-aware, and PaddleX splits a paragraph into multiple runs at every inline-formula /
    style boundary, so a compound can straddle two <w:t> nodes — converting per-node would mis-map
    a bisected phrase (公|里 -> 公裡 instead of 公里). So s2t runs at PARAGRAPH granularity: convert
    the joined text and redistribute by original node lengths (s2tw is 1:1; a rare length change
    falls back to per-node). All steps are idempotent (safe if a path applies them twice).
    """
    from docx import Document
    from docx.oxml.ns import qn

    from .i18n import normalize_locale

    is_tw = normalize_locale(locale) == "zh-TW"
    doc = Document(str(docx_path))
    # Body covers paragraphs + tables (they live in the body XML); headers/footers are
    # separate parts, so collect those too. Dedup shared parts by identity; header/footer
    # access is best-effort (suppressed).
    roots = [doc.element]
    for section in doc.sections:
        for hf in (section.header, section.footer):
            with contextlib.suppress(Exception):
                roots.append(hf.part.element)
    seen: set[int] = set()
    changed = False
    for root in roots:
        if id(root) in seen:
            continue
        seen.add(id(root))
        # Pass 1: per-<w:t> local normalization (entity decode, LaTeX render, stray Han spaces).
        for node in root.iter(qn("w:t")):
            if not node.text:
                continue
            fixed = html.unescape(node.text)  # `&amp;`/`&#8451;`->℃ that survive in native runs
            fixed = latex_to_unicode(fixed)
            fixed = strip_cjk_spaces(fixed)  # VL's stray Han<->Han spaces within a run
            if fixed != node.text:
                node.text = fixed
                changed = True
        # Pass 1.25: literal two-char "\n" inside TABLE cells -> real <w:br/> line breaks. VL
        # emits in-cell enumerated lists joined by a literal backslash-n; the markdown export
        # converts those to <br> (_fix_pipe_table_cell_newlines), so without this the two
        # exports of the same VL pass diverge and Word shows a verbatim "\n". Scoped to cells
        # (mirroring the markdown pass's pipe-row scoping) so a prose Windows path ("C:\name")
        # is never touched. The run is rebuilt as alternating <w:t>/<w:br/> children.
        from docx.oxml import OxmlElement

        for node in list(root.iter(qn("w:t"))):
            if not node.text or "\\n" not in node.text:
                continue
            if not any(True for _ in node.iterancestors(qn("w:tbl"))):
                continue
            run = node.getparent()
            if run is None or run.tag != qn("w:r"):
                continue
            parts = node.text.split("\\n")
            node.text = parts[0]
            node.set(qn("xml:space"), "preserve")
            anchor = node
            for part in parts[1:]:
                br = OxmlElement("w:br")
                anchor.addnext(br)
                t = OxmlElement("w:t")
                t.text = part
                t.set(qn("xml:space"), "preserve")
                br.addnext(t)
                anchor = t
            changed = True
        # Pass 1.5: Han<->Han spaces at RUN boundaries — invisible to the per-node strip above
        # (PaddleX puts "…繳交 " and "維護…" in separate <w:t> runs; real output shipped the
        # leaked space). Deletion-only, so the joined-string match spans map back exactly.
        for para in root.iter(qn("w:p")):
            nodes = [n for n in para.iter(qn("w:t")) if n.text]
            if len(nodes) < 2:
                continue
            stripped = strip_cjk_spaces_across([n.text for n in nodes])
            if stripped is not None:
                for n, t in zip(nodes, stripped, strict=True):
                    n.text = t
                changed = True
        # Pass 2 (zh-TW only): Simplified->Traditional per PARAGRAPH so a phrase split across runs
        # converts as a whole (see docstring). Redistribute by node length; fall back per-node if
        # the (near-1:1) conversion changed the length.
        if is_tw:
            for para in root.iter(qn("w:p")):
                nodes = [n for n in para.iter(qn("w:t")) if n.text]
                if not nodes:
                    continue
                joined = "".join(n.text for n in nodes)
                conv = fix_ocr_text(joined, locale)
                if conv == joined:
                    continue
                changed = True
                if len(conv) == len(joined):
                    pos = 0
                    for n in nodes:
                        n.text, pos = conv[pos : pos + len(n.text)], pos + len(n.text)
                else:
                    for n in nodes:
                        n.text = fix_ocr_text(n.text, locale)
    if changed:
        tmp = docx_path.with_suffix(".docx.tmp")
        doc.save(str(tmp))
        tmp.replace(docx_path)
