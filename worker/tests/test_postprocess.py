from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from worker.postprocess import (
    apply_word_heading_styles,
    apply_word_table_merges,
    drain_word_table_merge_specs,
    finalize_docx_text,
    html_img_to_markdown,
    html_table_to_markdown,
    process_markdown,
    reset_word_table_merge_specs,
    word_table_grid,
)


def _grid_texts(row: list[list[tuple[str, str]]]) -> list[str]:
    return ["".join(v for k, v in cell if k == "text") for cell in row]


def test_word_table_grid_rowspan_no_column_shift() -> None:
    # PaddleX's own DOCX table parser ignores rowspan and pads short rows at the END, shifting a
    # merged-cell continuation row one column left. Our rowspan-aware parser keeps a BLANK
    # leading cell so content stays in its column (fixes the docx-only cell-shift bug).
    html = (
        "<table>"
        "<tr><td>Cat</td><td>B</td><td>C</td></tr>"
        '<tr><td rowspan="2">Grp</td><td>b1</td><td>c1</td></tr>'
        "<tr><td>b2</td><td>c2</td></tr>"  # continuation row: only 2 <td> (Grp spans down)
        "</table>"
    )
    grid = word_table_grid(html)
    assert _grid_texts(grid[0]) == ["Cat", "B", "C"]
    assert _grid_texts(grid[1]) == ["Grp", "b1", "c1"]
    assert _grid_texts(grid[2]) == ["", "b2", "c2"]  # leading blank, NOT ["b2","c2",""]


def test_word_table_grid_colspan_expands_and_keeps_images() -> None:
    html = (
        "<table>"
        '<tr><td colspan="2">wide</td><td><img src="p.png"/></td></tr>'
        "<tr><td>a</td><td>b</td><td>c</td></tr>"
        "</table>"
    )
    grid = word_table_grid(html)
    assert _grid_texts(grid[0]) == ["wide", "", ""]  # colspan -> content then a blank column
    assert grid[0][2] == [("img", "p.png")]  # image segment preserved in the writer's shape
    assert _grid_texts(grid[1]) == ["a", "b", "c"]


def test_html_table_to_markdown() -> None:
    html = "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>"
    md = html_table_to_markdown(html)
    assert "| A | B |" in md
    assert "| --- | --- |" in md
    assert "| 1 | 2 |" in md


def test_html_table_colspan_pads_cells() -> None:
    html = '<table><tr><td colspan="2">wide</td></tr><tr><td>a</td><td>b</td></tr></table>'
    md = html_table_to_markdown(html)
    # colspan yields a 2-column table (empty trailing cell, not duplicated content)
    assert "| --- | --- |" in md
    assert "wide" in md
    assert "| a | b |" in md


def test_html_table_non_numeric_colspan_does_not_raise() -> None:
    # VL can emit annotated colspans like "2位" / "2.0"; a bare int() would raise ValueError
    # and fail the whole export. The leading integer is taken, so it still pads to 2 columns.
    html = '<table><tr><td colspan="2位">wide</td></tr><tr><td>a</td><td>b</td></tr></table>'
    md = html_table_to_markdown(html)
    assert "| --- | --- |" in md
    assert "wide" in md
    assert "| a | b |" in md


def test_html_table_garbage_colspan_defaults_to_one() -> None:
    # A colspan with no leading digits must not crash; it falls back to 1.
    html = '<table><tr><td colspan="abc">x</td></tr></table>'
    md = html_table_to_markdown(html)
    assert "x" in md


def test_html_table_giant_colspan_is_clamped() -> None:
    # A garbage/huge colspan must not balloon the row to thousands of columns (memory blowup);
    # it is clamped to _MAX_SPAN (64). Normal small spans are unaffected.
    html = '<table><tr><td colspan="9999">wide</td></tr><tr><td>a</td></tr></table>'
    md = html_table_to_markdown(html)
    header = next(ln for ln in md.splitlines() if ln.startswith("| wide"))
    assert header.count("|") == 65  # 64 clamped columns -> 65 pipes, not ~10000


def test_html_table_rowspan_keeps_columns_aligned() -> None:
    # A rowspan cell must not shift the cells below it left. Markdown has no rowspan, so the
    # spanned-down position is emitted blank and the remaining cells stay column-aligned.
    html = (
        "<table>"
        "<tr><th>Category</th><th>Item</th><th>Amount</th></tr>"
        '<tr><td rowspan="2">Income</td><td>Salary</td><td>5000</td></tr>'
        "<tr><td>Bonus</td><td>1000</td></tr>"
        "</table>"
    )
    lines = [ln for ln in html_table_to_markdown(html).splitlines() if ln.startswith("|")]
    assert lines[2] == "| Income | Salary | 5000 |"
    assert lines[3] == "|   | Bonus | 1000 |"  # Category held by rowspan -> blank, not shifted


def test_process_markdown_unwraps_vl_center_div_caption() -> None:
    # VL wraps figure/table captions in a centered div; markdown has no div, so unwrap it.
    md = process_markdown('<div style="text-align: center;">Figure 1: results</div>', None, "en", images=False)
    assert "<div" not in md
    assert "Figure 1: results" in md


def test_process_markdown_preserves_real_div_in_body_text() -> None:
    # A genuine <div> in OCR'd body text (not the exact VL center wrapper) must survive.
    src = 'Write <div class="box"> then </div> to close.'
    md = process_markdown(src, None, "en", images=False)
    assert '<div class="box">' in md


def test_html_img_generic_alt_is_stripped() -> None:
    md = html_img_to_markdown('<img alt="Image" src="pic.png"/>', None)
    assert md.strip() == "![](pic.png)"


def test_process_markdown_converts_table_and_fixes_alt() -> None:
    src = '<table><tr><td>x</td></tr></table>\n<img alt="Image" src="a.png"/>'
    md = process_markdown(src, None, "en", images=True)
    assert "| x |" in md
    assert "![](a.png)" in md


def test_process_markdown_decodes_html_entities_in_prose() -> None:
    # VL leaves HTML entities in prose outside tables; they must ship decoded, not literal.
    md = process_markdown("研究 A &amp; B 在 &#8451; 與 &deg;C 下", None, "en", images=False)
    assert "A & B" in md and "℃" in md and "°C" in md
    assert "&amp;" not in md and "&#8451;" not in md


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    return "".join(run.text for para in doc.paragraphs for run in para.runs)


def test_finalize_docx_converts_simplified_for_zh_tw(tmp_path: Path) -> None:
    # The native save_to_word path (PaddleOCR-VL) bypasses process_markdown; without the
    # docx-level pass a zh-TW Word export keeps VL's Simplified output — including the
    # ambiguous chars the old static table skipped (范围 -> 範圍).
    p = tmp_path / "t.docx"
    d = Document()
    d.add_paragraph("系统数据范围")  # 系统->系統 数据->數據 范围->範圍 (OpenCC)
    d.save(str(p))
    finalize_docx_text(p, "zh-TW")
    text = _docx_text(p)
    assert "系統數據範圍" in text
    assert "统" not in text and "范围" not in text


def test_finalize_docx_converts_phrase_split_across_runs(tmp_path: Path) -> None:
    # OpenCC s2tw is phrase-aware, not a char map: PaddleX splits a paragraph into multiple runs
    # at formula/style boundaries, so a Simplified compound can straddle two runs (<w:t> nodes).
    # Per-node conversion mis-maps it ("公"+"里..." -> 公裡); paragraph-level conversion yields 公里.
    p = tmp_path / "t.docx"
    d = Document()
    para = d.add_paragraph()
    para.add_run("公")  # 公里 bisected across two runs...
    para.add_run("里与头发")  # ...second half + more (与->與, 头发->頭髮)
    d.save(str(p))
    finalize_docx_text(p, "zh-TW")
    runs = Document(str(p)).paragraphs[0].runs
    assert "".join(r.text for r in runs) == "公里與頭髮"  # NOT 公裡 (the per-node bug)
    assert [r.text for r in runs] == ["公", "里與頭髮"]  # run boundaries preserved (redistribute by len)


def test_finalize_docx_strips_cjk_space_at_run_boundary(tmp_path: Path) -> None:
    # A stray Han<->Han space can sit exactly at a run boundary ("…繳交 " | "維護…" — PaddleX
    # splits runs at style changes), where the per-node strip can't see it. Real 06-doc output
    # shipped that space. The cross-run pass must drop it — while a legit CJK<->Latin boundary
    # space ("通知後 24") survives.
    p = tmp_path / "t.docx"
    d = Document()
    para = d.add_paragraph()
    para.add_run("日曆天內繳交 ")
    para.add_run("維護人員名冊")
    para2 = d.add_paragraph()
    para2.add_run("通知後 ")  # Han + space, but the next run starts with a digit -> must survive
    para2.add_run("24小時以內")
    d.save(str(p))
    finalize_docx_text(p, "zh-TW")
    doc = Document(str(p))
    assert "".join(r.text for r in doc.paragraphs[0].runs) == "日曆天內繳交維護人員名冊"
    assert "".join(r.text for r in doc.paragraphs[1].runs) == "通知後 24小時以內"


def test_apply_word_heading_styles_uses_markdown_levels(tmp_path: Path) -> None:
    # PaddleX styles EVERY heading level at the same ~14pt bold, so size-ranking flattens the
    # outline to two levels. A levels map harvested from the same VL pass's markdown must set the
    # TRUE depth for text-matched headings; unmatched ones keep the size-rank fallback.
    from docx.shared import Pt

    p = tmp_path / "t.docx"
    d = Document()
    for text in ("壹、總則", "一、目的", "（一）範圍"):
        para = d.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(14)  # all the same size — rank alone cannot recover depth
    for _ in range(12):
        d.add_paragraph("內文段落，不是標題。")  # body: not bold -> Normal
    d.save(str(p))
    levels = {"壹、總則": 2, "一、目的": 3, "（一）範圍": 4}
    apply_word_heading_styles(p, levels)
    doc = Document(str(p))

    def _style(i: int) -> str:
        s = doc.paragraphs[i].style
        return s.name or "" if s is not None else ""

    assert [_style(i) for i in range(3)] == ["Heading 2", "Heading 3", "Heading 4"]
    assert _style(3) == "Normal"


def test_finalize_docx_converts_literal_newline_in_table_cells(tmp_path: Path) -> None:
    # VL joins in-cell enumerated lists with a literal two-char "\n"; the markdown export turns
    # those into <br>, so the DOCX must turn them into real <w:br/> breaks — otherwise Word
    # renders a verbatim backslash-n and the two exports of one VL pass diverge. Prose (incl. a
    # Windows path) is never touched: the pass is scoped to table cells.
    from docx.oxml.ns import qn

    p = tmp_path / "t.docx"
    d = Document()
    d.add_paragraph("path C:\\network stays")  # prose: contains literal \n ("\\n" in "\\network")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "1. 完成安裝\\n2. 完成上傳\\n3. 完成測試"
    d.save(str(p))
    finalize_docx_text(p, "zh-TW")
    doc = Document(str(p))
    tbl = doc.tables[0]._tbl
    assert len(tbl.findall(".//" + qn("w:br"))) == 2
    cell_text = "".join(n.text or "" for n in tbl.iter(qn("w:t")))
    assert "\\n" not in cell_text
    assert "完成安裝" in cell_text and "完成測試" in cell_text
    assert "C:\\network" in doc.paragraphs[0].text  # prose untouched


def test_finalize_docx_renders_latex_all_locales(tmp_path: Path) -> None:
    # Inline LaTeX must render to Unicode regardless of locale (VL emits it for units/math).
    for loc in ("zh-TW", "zh-CN", "en"):
        p = tmp_path / f"t_{loc}.docx"
        d = Document()
        d.add_paragraph("量測精度 $ \\pm0.3 $ $ ^{\\circ} $C，$ 8mm^{2} $ 導線。")
        d.save(str(p))
        finalize_docx_text(p, loc)
        text = _docx_text(p)
        assert "±0.3" in text and "°C" in text and "8mm²" in text
        assert "$" not in text and "\\circ" not in text


def test_finalize_docx_decodes_html_entities(tmp_path: Path) -> None:
    # Native save_to_word runs can carry HTML entities; finalize must decode them (all locales).
    p = tmp_path / "t.docx"
    d = Document()
    d.add_paragraph("A &amp; B 於 &#8451;")
    d.save(str(p))
    finalize_docx_text(p, "en")
    text = _docx_text(p)
    assert "A & B" in text and "℃" in text
    assert "&amp;" not in text and "&#8451;" not in text


def test_finalize_docx_noop_for_non_zh_tw_leaves_simplified(tmp_path: Path) -> None:
    # zh-CN / en users keep native Simplified output — only LaTeX is normalized, not script.
    p = tmp_path / "t.docx"
    d = Document()
    d.add_paragraph("系统统计")  # no LaTeX, non-zh-TW -> unchanged
    d.save(str(p))
    finalize_docx_text(p, "zh-CN")
    assert "统" in _docx_text(p)


# --- #2: merged-cell (rowspan/colspan) -> real w:vMerge / w:gridSpan --------------------------
def _grid_to_docx(grid: list[list[list[tuple[str, str]]]], path: Path) -> None:
    # Fill a "Table Grid" table cell-by-cell exactly like PaddleX's writer (word_converter.py:583),
    # so the merge pass runs against the real emitted shape (one <w:tc> per grid cell, no merges).
    from docx import Document

    doc = Document()
    max_cols = max(len(r) for r in grid)
    t = doc.add_table(rows=0, cols=max_cols)
    t.style = "Table Grid"
    for row_cells in grid:
        cells = t.add_row().cells
        for i in range(max_cols):
            segs = row_cells[i] if i < len(row_cells) else []
            cells[i].text = "".join(v for k, v in segs if k == "text")
    doc.save(str(path))


def _tbl_row_cell_texts(tbl: Any) -> list[list[str]]:
    from docx.oxml.ns import qn

    rows = []
    for tr in tbl.findall(qn("w:tr")):
        rows.append(["".join(n.text or "" for n in tc.iter(qn("w:t"))) for tc in tr.findall(qn("w:tc"))])
    return rows


def test_word_table_grid_records_span_specs() -> None:
    # word_table_grid appends one MergeSpec per table: {(row,col): (rowspan,colspan)} for span
    # ORIGINS only, indexing the SAME rectangular grid it returns.
    reset_word_table_merge_specs()
    html = (
        "<table>"
        "<tr><td>H</td><td>a</td><td>b</td><td>c</td></tr>"
        '<tr><td rowspan="2">grp</td><td>d</td><td>e</td><td>f</td></tr>'
        "<tr><td>g</td><td>h</td><td>i</td></tr>"
        '<tr><td colspan="4">wide</td></tr>'
        "</table>"
    )
    word_table_grid(html)
    specs = drain_word_table_merge_specs()
    # One entry per table: (span origins, row count, first-row cell count, first-row fingerprint).
    assert specs == [({(1, 0): (2, 1), (3, 0): (1, 4)}, 4, 4, "Habc")]
    assert drain_word_table_merge_specs() == []  # drained -> empty


def test_apply_word_table_merges_writes_vmerge_and_gridspan(tmp_path: Path) -> None:
    from docx import Document

    reset_word_table_merge_specs()
    html = (
        "<table>"
        "<tr><td>H</td><td>a</td><td>b</td><td>c</td></tr>"
        '<tr><td rowspan="2">grp</td><td>d</td><td>e</td><td>f</td></tr>'
        "<tr><td>g</td><td>h</td><td>i</td></tr>"
        '<tr><td colspan="4">wide</td></tr>'
        "</table>"
    )
    grid = word_table_grid(html)
    specs = drain_word_table_merge_specs()
    p = tmp_path / "t.docx"
    _grid_to_docx(grid, p)
    apply_word_table_merges(p, specs)

    from docx.oxml.ns import qn

    tbl = Document(str(p)).tables[0]._tbl
    vmerges = tbl.findall(".//" + qn("w:vMerge"))
    gridspans = tbl.findall(".//" + qn("w:gridSpan"))
    # rowspan=2 -> one restart + one continue (val absent) vMerge
    assert sorted((e.get(qn("w:val")) or "continue") for e in vmerges) == ["continue", "restart"]
    # colspan=4 -> one gridSpan of 4, and the 3 blank continuation cells removed from that row
    assert [e.get(qn("w:val")) for e in gridspans] == ["4"]
    texts = _tbl_row_cell_texts(tbl)
    assert texts[1] == ["grp", "d", "e", "f"]  # rowspan origin cell kept
    assert texts[2] == ["", "g", "h", "i"]  # continuation cell kept (blank, marked vMerge)
    assert texts[3] == ["wide"]  # colspan collapsed to a single gridSpan cell


def test_apply_word_table_merges_noop_for_unmerged_table(tmp_path: Path) -> None:
    # An empty spec (no rowspan/colspan anywhere) must leave the DOCX byte-identical.
    grid = [[[("text", "a")], [("text", "b")]], [[("text", "c")], [("text", "d")]]]
    p = tmp_path / "t.docx"
    _grid_to_docx(grid, p)
    before = p.read_bytes()
    apply_word_table_merges(p, [({}, 2, 2, "ab")])
    assert p.read_bytes() == before


def test_apply_word_table_merges_skips_on_shape_mismatch(tmp_path: Path) -> None:
    # A spec whose coords fall outside the realized table (e.g. merge_docx reordering) must be
    # skipped gracefully, never corrupt: no vMerge/gridSpan emitted, cells intact.
    from docx import Document
    from docx.oxml.ns import qn

    grid = [[[("text", "a")], [("text", "b")]]]  # 1x2 table
    p = tmp_path / "t.docx"
    _grid_to_docx(grid, p)
    apply_word_table_merges(p, [({(9, 9): (3, 3)}, 1, 2, "ab")])  # out-of-range coords
    tbl = Document(str(p)).tables[0]._tbl
    assert tbl.findall(".//" + qn("w:vMerge")) == []
    assert tbl.findall(".//" + qn("w:gridSpan")) == []
    assert _tbl_row_cell_texts(tbl) == [["a", "b"]]


def test_apply_word_table_merges_matches_by_fingerprint_not_position(tmp_path: Path) -> None:
    # PaddleX's writer can emit <w:tbl> for NON-HTML tables the collector never saw; a spec must
    # find ITS table by fingerprint, not land positionally on the first table in the body.
    from docx import Document
    from docx.oxml.ns import qn

    reset_word_table_merge_specs()
    html = (
        "<table>"
        '<tr><td rowspan="2">grp</td><td>x</td></tr>'
        "<tr><td>y</td></tr>"
        "</table>"
    )
    grid = word_table_grid(html)
    specs = drain_word_table_merge_specs()
    p = tmp_path / "t.docx"
    # Body contains an unrelated (non-HTML) 2x2 table FIRST, then the HTML table the spec is for.
    other = [[[("text", "p")], [("text", "q")]], [[("text", "r")], [("text", "s")]]]
    _grid_to_docx(other, p)
    d = Document(str(p))
    t = d.add_table(rows=len(grid), cols=2)
    for ri, row in enumerate(grid):
        for ci, cell in enumerate(row):
            t.cell(ri, ci).text = "".join(v for k, v in cell if k == "text")
    d.save(str(p))
    apply_word_table_merges(p, specs)
    doc = Document(str(p))
    first, second = doc.tables[0]._tbl, doc.tables[1]._tbl
    assert first.findall(".//" + qn("w:vMerge")) == []  # untouched: not the spec's table
    assert len(second.findall(".//" + qn("w:vMerge"))) == 2  # restart + continue landed here


# --- #1: literal "\n" inside a native GFM pipe-table cell -> "<br>" ---------------------------
def test_process_markdown_converts_literal_newline_in_pipe_cell() -> None:
    src = "| 驗收標準 | x |\n| --- | --- |\n| 1. 完成安裝\\n2. 完成上傳 | ok |"
    md = process_markdown(src, None, "en", images=False)
    assert "1. 完成安裝<br>2. 完成上傳" in md
    assert "\\n" not in md


def test_process_markdown_keeps_literal_newline_outside_tables() -> None:
    # A literal backslash-n in prose (a Windows path, a code sample) must NOT become <br>.
    src = "Open the file at C:\\name and continue.\n\nA line with \\n in text."
    md = process_markdown(src, None, "en", images=False)
    assert "C:\\name" in md
    assert "\\n in text" in md
    assert "<br>" not in md


def test_process_markdown_pipe_cell_newline_is_idempotent() -> None:
    src = "| a | b |\n| --- | --- |\n| 1\\n2 | y |"
    once = process_markdown(src, None, "en", images=False)
    twice = process_markdown(once, None, "en", images=False)
    assert once == twice


# --- #8: size/bold heading paragraphs -> real 'Heading N' styles ------------------------------
def _style_name(para: Any) -> str | None:
    return para.style.name if para.style is not None else None


def test_apply_word_heading_styles_ranks_sizes_to_levels(tmp_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    d = Document()

    def _head(text: str, size: int) -> None:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = True

    _head("Chapter", 20)  # largest -> Heading 1
    _head("Section", 14)  # next -> Heading 2
    for _ in range(6):  # body dominates so the heading fraction stays well under the guard
        body = d.add_paragraph()
        body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = body.add_run("body text here")
        run.font.size = Pt(12)
    path = tmp_path / "t.docx"
    d.save(str(path))
    apply_word_heading_styles(path)

    styles = [_style_name(para) for para in Document(str(path)).paragraphs if para.text.strip()]
    assert styles[0] == "Heading 1"
    assert styles[1] == "Heading 2"
    assert all(s == "Normal" for s in styles[2:])  # justified body untouched


def test_apply_word_heading_styles_bails_when_everything_looks_like_heading(tmp_path: Path) -> None:
    # If more than half the paragraphs are bold+big, "heading" is not a real signal -> leave as-is.
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    d = Document()
    for _ in range(5):
        para = d.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run("all bold big")
        run.font.size = Pt(14)
        run.font.bold = True
    path = tmp_path / "t.docx"
    d.save(str(path))
    apply_word_heading_styles(path)
    assert all(_style_name(para) == "Normal" for para in Document(str(path)).paragraphs if para.text.strip())
