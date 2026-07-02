from pathlib import Path

import pytest

# doc2md needs paddleocr installed; skip where it isn't (e.g. CI without the GPU stack).
pytest.importorskip("paddleocr")

from worker.office import is_office, process
from worker.store import Store

SCHEMA = Path(__file__).resolve().parents[2] / "db" / "schema.sql"


def test_is_office() -> None:
    assert is_office("/x/a.docx")
    assert is_office("/x/b.pptx")
    assert is_office("/x/c.xlsx")
    assert not is_office("/x/d.pdf")


def test_docx_to_markdown(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from docx import Document

    import worker.config as cfg

    monkeypatch.setattr(cfg, "OUTPUTS_DIR", tmp_path / "out")

    docx = tmp_path / "t.docx"
    doc = Document()
    doc.add_heading("標題 Title", 0)
    doc.add_paragraph("內容 body")
    doc.save(docx)

    store = Store(tmp_path / "t.db", SCHEMA)
    store.conn.execute(
        "INSERT INTO jobs(id,mode,filename,locale,status,upload_path,created_at,updated_at) "
        "VALUES('o1','markdown','t.docx','en','processing',?,1,1)",
        (str(docx),),
    )
    process(store, {"id": "o1", "locale": "en", "filename": "t.docx", "upload_path": str(docx)})

    md = (tmp_path / "out" / "o1" / "o1.md").read_text(encoding="utf-8")
    assert "標題 Title" in md
    assert "內容 body" in md
    assert store.status_of("o1") == "done"
